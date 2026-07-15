from __future__ import annotations

import base64
import io
import json
import mimetypes
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

from app.config import Config
from steeltech_db.extensions import db
from steeltech_db.models.weekly_meeting import (
    WeeklyMeeting,
    WeeklyMeetingImage,
    WeeklyMeetingScan,
)

WEEKLY_MEETING_IMAGE_ROOT = Path(Config.WEEKLY_MEETING_IMAGE_ROOT)
WEEKLY_MEETING_SCAN_ROOT = Path(Config.WEEKLY_MEETING_SCAN_ROOT)
MAX_MEETING_IMAGES = 2


def _now_local() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _generate_id() -> str:
    return uuid.uuid4().hex[:16]


def _meeting_to_dict(meeting: WeeklyMeeting) -> dict:
    images = (
        WeeklyMeetingImage.query
        .filter_by(meeting_id=meeting.id)
        .order_by(WeeklyMeetingImage.sort_order)
        .all()
    )
    scan = WeeklyMeetingScan.query.filter_by(meeting_id=meeting.id).first()
    return {
        "id": meeting.id,
        "meetingDate": meeting.meeting_date,
        "meetingTheme": meeting.meeting_theme,
        "location": meeting.location,
        "host": meeting.host,
        "recorder": meeting.recorder,
        "attendees": meeting.attendees,
        "contentItems": json.loads(meeting.content_items or "[]"),
        "images": [_image_to_dict(img) for img in images],
        "recordScan": _scan_to_dict(scan) if scan else None,
        "createdAt": meeting.created_at,
        "updatedAt": meeting.updated_at,
    }


def _image_to_dict(image: WeeklyMeetingImage) -> dict:
    return {
        "id": image.id,
        "meetingId": image.meeting_id,
        "fileName": image.file_name,
        "filePath": image.file_path,
        "fileSize": image.file_size,
        "mimeType": image.mime_type,
        "sortOrder": image.sort_order,
        "url": f"/api/weekly-meeting-images/{image.file_path}",
    }


def _scan_to_dict(scan: WeeklyMeetingScan) -> dict:
    return {
        "id": scan.id,
        "meetingId": scan.meeting_id,
        "fileName": scan.file_name,
        "filePath": scan.file_path,
        "fileSize": scan.file_size,
        "mimeType": scan.mime_type,
        "url": f"/api/weekly-meeting-scans/{scan.file_path}",
    }


# ── CRUD ────────────────────────────────────────────────────────────────────


def list_meetings(year: str = "", month: str = "") -> dict:
    query = WeeklyMeeting.query
    if year:
        query = query.filter(WeeklyMeeting.meeting_date.like(f"{year}-%"))
    if month:
        query = query.filter(WeeklyMeeting.meeting_date.like(f"%-{month}-%"))
    query = query.order_by(WeeklyMeeting.meeting_date.desc(), WeeklyMeeting.id.desc())

    meetings = query.all()

    return {
        "list": [_meeting_to_dict(m) for m in meetings],
        "total": len(meetings),
    }


def get_meeting(meeting_id: str) -> dict | None:
    meeting = WeeklyMeeting.query.get(meeting_id)
    if not meeting:
        return None
    return _meeting_to_dict(meeting)


def create_meeting(payload: dict) -> dict:
    meeting = WeeklyMeeting(
        id=_generate_id(),
        meeting_date=str(payload.get("meetingDate", "")).strip(),
        meeting_theme=str(payload.get("meetingTheme", "")).strip(),
        location=str(payload.get("location", "")).strip(),
        host=str(payload.get("host", "")).strip(),
        recorder=str(payload.get("recorder", "")).strip(),
        attendees=str(payload.get("attendees", "")).strip(),
        content_items=json.dumps(payload.get("contentItems", []), ensure_ascii=False),
        created_at=_now_local(),
    )
    db.session.add(meeting)
    db.session.commit()
    return _meeting_to_dict(meeting)


def update_meeting(meeting_id: str, payload: dict) -> dict | None:
    meeting = WeeklyMeeting.query.get(meeting_id)
    if not meeting:
        return None

    if "meetingDate" in payload:
        meeting.meeting_date = str(payload["meetingDate"]).strip()
    if "meetingTheme" in payload:
        meeting.meeting_theme = str(payload["meetingTheme"]).strip()
    if "location" in payload:
        meeting.location = str(payload["location"]).strip()
    if "host" in payload:
        meeting.host = str(payload["host"]).strip()
    if "recorder" in payload:
        meeting.recorder = str(payload["recorder"]).strip()
    if "attendees" in payload:
        meeting.attendees = str(payload["attendees"]).strip()
    if "contentItems" in payload:
        meeting.content_items = json.dumps(payload["contentItems"], ensure_ascii=False)

    meeting.updated_at = _now_local()
    db.session.commit()
    return _meeting_to_dict(meeting)


def delete_meeting(meeting_id: str) -> bool:
    meeting = WeeklyMeeting.query.get(meeting_id)
    if not meeting:
        return False

    images = WeeklyMeetingImage.query.filter_by(meeting_id=meeting_id).all()
    for img in images:
        _delete_image_file(img)
        db.session.delete(img)

    scan = WeeklyMeetingScan.query.filter_by(meeting_id=meeting_id).first()
    if scan:
        _delete_scan_file(scan)
        db.session.delete(scan)

    db.session.delete(meeting)
    db.session.commit()
    return True


# ── Image management ────────────────────────────────────────────────────────


def _delete_image_file(image: WeeklyMeetingImage) -> None:
    try:
        file_path = WEEKLY_MEETING_IMAGE_ROOT / image.file_path
        if file_path.exists():
            file_path.unlink()
    except Exception:
        pass


def save_meeting_images(meeting_id: str, files: list[dict]) -> list[dict]:
    """保存上传的图片。files 每项: {fileName, content(base64), mimeType}"""
    meeting = WeeklyMeeting.query.get(meeting_id)
    if not meeting:
        raise ValueError("会议记录不存在")

    existing_count = WeeklyMeetingImage.query.filter_by(meeting_id=meeting_id).count()
    if existing_count + len(files) > MAX_MEETING_IMAGES:
        raise ValueError(f"每条会议记录最多上传 {MAX_MEETING_IMAGES} 张图片")

    current_max_order = (
        db.session.query(db.func.max(WeeklyMeetingImage.sort_order))
        .filter_by(meeting_id=meeting_id)
        .scalar()
    ) or 0

    saved = []
    meeting_dir = WEEKLY_MEETING_IMAGE_ROOT / meeting_id
    meeting_dir.mkdir(parents=True, exist_ok=True)

    for index, file_info in enumerate(files):
        file_name = str(file_info.get("fileName", "image.png"))
        content_b64 = str(file_info.get("content", ""))
        mime_type = str(file_info.get("mimeType", "image/png"))

        ext = _guess_ext(file_name, mime_type)
        stored_name = f"{uuid.uuid4().hex[:8]}{ext}"
        relative_path = f"{meeting_id}/{stored_name}"
        full_path = meeting_dir / stored_name

        try:
            raw_bytes = base64.b64decode(content_b64)
        except Exception:
            continue

        full_path.write_bytes(raw_bytes)

        image = WeeklyMeetingImage(
            id=_generate_id(),
            meeting_id=meeting_id,
            file_name=file_name,
            file_path=relative_path,
            file_size=len(raw_bytes),
            mime_type=mime_type,
            sort_order=current_max_order + index + 1,
            created_at=_now_local(),
        )
        db.session.add(image)
        saved.append(image)

    db.session.commit()
    return [_image_to_dict(img) for img in saved]


def delete_meeting_image(meeting_id: str, image_id: str) -> bool:
    image = WeeklyMeetingImage.query.filter_by(
        id=image_id, meeting_id=meeting_id
    ).first()
    if not image:
        return False
    _delete_image_file(image)
    db.session.delete(image)
    db.session.commit()
    return True


def _guess_ext(file_name: str, mime_type: str) -> str:
    if mime_type:
        ext = mimetypes.guess_extension(mime_type)
        if ext:
            return ext
    return Path(file_name).suffix or ".png"


# ── Scan management ─────────────────────────────────────────────────────────


def _delete_scan_file(scan: WeeklyMeetingScan) -> None:
    try:
        file_path = WEEKLY_MEETING_SCAN_ROOT / scan.file_path
        if file_path.exists():
            file_path.unlink()
    except Exception:
        pass


def _is_pdf_file(file_name: str, mime_type: str) -> bool:
    normalized_name = file_name.lower()
    normalized_mime = mime_type.lower()
    return normalized_name.endswith(".pdf") or normalized_mime == "application/pdf"


def save_meeting_scan(meeting_id: str, file_info: dict) -> dict:
    """保存会议记录扫描件 PDF。每条会议记录仅保留一份，重复上传会覆盖。"""
    meeting = WeeklyMeeting.query.get(meeting_id)
    if not meeting:
        raise ValueError("会议记录不存在")

    file_name = str(file_info.get("fileName", "scan.pdf")).strip()
    content_b64 = str(file_info.get("content", "")).strip()
    mime_type = str(file_info.get("mimeType", "application/pdf")).strip()

    if not content_b64:
        raise ValueError("没有上传文件")

    if not _is_pdf_file(file_name, mime_type):
        raise ValueError("会议记录扫描件仅支持 PDF 格式")

    try:
        raw_bytes = base64.b64decode(content_b64)
    except Exception as exc:
        raise ValueError("文件内容无效") from exc

    if not raw_bytes:
        raise ValueError("文件内容为空")

    existing = WeeklyMeetingScan.query.filter_by(meeting_id=meeting_id).first()
    if existing:
        _delete_scan_file(existing)
        db.session.delete(existing)
        db.session.flush()

    meeting_dir = WEEKLY_MEETING_SCAN_ROOT / meeting_id
    meeting_dir.mkdir(parents=True, exist_ok=True)

    stored_name = f"{uuid.uuid4().hex[:8]}.pdf"
    relative_path = f"{meeting_id}/{stored_name}"
    full_path = meeting_dir / stored_name
    full_path.write_bytes(raw_bytes)

    scan = WeeklyMeetingScan(
        id=_generate_id(),
        meeting_id=meeting_id,
        file_name=file_name,
        file_path=relative_path,
        file_size=len(raw_bytes),
        mime_type="application/pdf",
        created_at=_now_local(),
    )
    db.session.add(scan)
    db.session.commit()
    return _scan_to_dict(scan)


def delete_meeting_scan(meeting_id: str) -> bool:
    scan = WeeklyMeetingScan.query.filter_by(meeting_id=meeting_id).first()
    if not scan:
        return False
    _delete_scan_file(scan)
    db.session.delete(scan)
    db.session.commit()
    return True


# ── DOCX Export ─────────────────────────────────────────────────────────────


def export_meeting_docx(meeting_id: str) -> bytes | None:
    """生成周例会 DOCX 文件（含 Lampiran A 签到表 + Lampiran B 纪要）"""
    meeting_data = get_meeting(meeting_id)
    if not meeting_data:
        return None

    # Lazy import — only when actually exporting
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")

    # ── Nested helpers (closure avoids any global dependency) ────────────

    def _cn_weekday(date_str: str) -> str:
        try:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            return ["周一","周二","周三","周四","周五","周六","周日"][dt.weekday()]
        except Exception:
            return ""

    def _set_run_east_asian(run, font_name: str):
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _set_cell_text(table, row: int, col: int, text: str, bold: bool = False, fs: float = 10.5):
        cell = table.cell(row, col)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(fs)
        run.font.name = "Times New Roman"
        return run

    def _merge_cells(table, rs: int, cs: int, re: int, ce: int):
        table.cell(rs, cs).merge(table.cell(re, ce))

    def _add_page_break():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _set_row_height(row, height_cm: float):
        row.height = Cm(height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # ── Parse data ────────────────────────────────────────────────────────
    meeting_date = meeting_data["meetingDate"]
    meeting_theme = meeting_data["meetingTheme"]
    location = meeting_data["location"]
    host = meeting_data["host"]
    recorder = meeting_data["recorder"]
    attendees = meeting_data["attendees"]
    content_items = meeting_data.get("contentItems", [])

    # 签到表姓名列表：主持排第一，参会人员按逗号分隔排后面
    attendee_names = [name.strip() for name in attendees.split(",") if name.strip()] if attendees else []
    host_name = host.strip() if host else ""
    sign_in_names: list[str] = []
    if host_name:
        sign_in_names.append(host_name)
    for name in attendee_names:
        if name and name != host_name:
            sign_in_names.append(name)

    try:
        dt = datetime.strptime(meeting_date, "%Y-%m-%d")
        year_str = str(dt.year)
        month_str = f"{dt.month:02d}"
        day_str = f"{dt.day:02d}"
        date_display = f"{dt.year}.{dt.month}.{dt.day}"
    except Exception:
        dt = datetime(2026, 6, 12)
        year_str, month_str, day_str = "2026", "06", "12"
        date_display = meeting_date

    weekday = _cn_weekday(meeting_date)

    # ── Build document ────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ===== Lampiran A: Sign-in Sheet (独占第一页) ===========================

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Lampiran A")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(
        "Daftar Hadir Rapat PT ITSS Departemen Teknologi Peralatan\n"
        "PT ITSS 设备技术部会议签到表"
    )
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"

    # 签到行数：至少 10 行，根据实际人数增加
    sign_in_count = max(10, len(sign_in_names))
    total_rows = 4 + sign_in_count  # 4 header rows + sign-in rows
    table = doc.add_table(rows=total_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(3.2)
        row.cells[3].width = Cm(4.5)

    _set_cell_text(table, 0, 0, "Tema Rapat\n会议主题", bold=True)
    _set_cell_text(table, 0, 1, meeting_theme)
    _merge_cells(table, 0, 1, 0, 3)

    _set_cell_text(table, 1, 0, "Waktu Rapat\n会议时间", bold=True)
    _set_cell_text(table, 1, 1, date_display)
    _set_cell_text(table, 1, 2, "Lokasi Rapat\n会议地点", bold=True)
    _set_cell_text(table, 1, 3, location)

    _set_cell_text(table, 2, 0, "Moderator\n主持", bold=True)
    _set_cell_text(table, 2, 1, host)
    _set_cell_text(table, 2, 2, "Peserta Rapat\n参会人员", bold=True)
    _set_cell_text(table, 2, 3, attendees)

    _set_cell_text(table, 3, 0, "No\n序号", bold=True)
    _set_cell_text(table, 3, 1, "Departemen\n部门", bold=True)
    _set_cell_text(table, 3, 2, "Nama\n姓名", bold=True)
    _set_cell_text(table, 3, 3, "Tanda Tangan\n签名", bold=True)

    for i in range(sign_in_count):
        _set_cell_text(table, 4 + i, 0, str(i + 1))
        if i < len(sign_in_names):
            _set_cell_text(table, 4 + i, 2, sign_in_names[i])

    for row_index, row in enumerate(table.rows):
        if row_index <= 3:
            _set_row_height(row, 1.15)
        else:
            _set_row_height(row, 1.58)

    _add_page_break()

    # ===== Lampiran B: Meeting Minutes (独占第二页) =========================

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lampiran B")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Notulen Rapat Mingguan\n周例会纪要")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    def _add_field(label_id: str, label_cn: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(22)
        run = p.add_run(f"{label_id}: ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run = p.add_run(f"{label_cn}: ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        _set_run_east_asian(run, "宋体")
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    _add_field("Tema Rapat", "会议主题", f"{meeting_theme} / 钢结构技术科周例会")

    weekday_str = weekday or ""
    _add_field("Waktu Rapat", "会议时间",
        f"{year_str} Tahun {month_str} Bulan {day_str} Tanggal "
        f"14 Jam 00 Menit Sampai 15 Jam 00 Menit\n"
        f"{year_str} 年 {month_str} 月 {day_str} 日 14:00 - 15:00 {weekday_str}")

    _add_field("Tempat Rapat", "会议地点", location)
    _add_field("Pemimpin Rapat", "会议主持", host)
    _add_field("Catatan Rapat", "会议记录", recorder)
    _add_field("Peserta", "出席人员", attendees)

    # Content items
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run("Konten Rapat: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run = p.add_run("会议内容: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    _set_run_east_asian(run, "宋体")

    if content_items:
        for i, item in enumerate(content_items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(22)
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"{i}. {str(item or '').strip()}")
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
    else:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("(会议内容待填写)")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # Sign-off table
    doc.add_paragraph()
    doc.add_paragraph()

    sign_table = doc.add_table(rows=2, cols=3)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_labels = [
        ["Dibuat / 编制", "Diperiksa / 审核", "Disetujui / 批准"],
        ["(Tanda Tangan / 签名)", "(Tanda Tangan / 签名)", "(Tanda Tangan / 签名)"],
    ]
    for ri in range(2):
        for ci in range(3):
            cell = sign_table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(sign_labels[ri][ci])
            if ri == 0:
                run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"

    # ===== 会议图片页 (独占最后一页) =========================================

    meeting_images = meeting_data.get("images", [])[:MAX_MEETING_IMAGES]
    if meeting_images:
        _add_page_break()

        for image_info in meeting_images:
            image_path = WEEKLY_MEETING_IMAGE_ROOT / image_info["filePath"]
            if not image_path.exists():
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            run.add_picture(str(image_path), width=Cm(15.5))

    # ── Save to bytes ────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
