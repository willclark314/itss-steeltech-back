from __future__ import annotations

import base64
import io
import json
import mimetypes
import uuid
from datetime import datetime
from pathlib import Path

from app.config import Config
from steeltech_db.extensions import db
from steeltech_db.models.director_interview import DirectorInterview, DirectorInterviewImage

DIRECTOR_INTERVIEW_IMAGE_ROOT = Path(Config.DIRECTOR_INTERVIEW_IMAGE_ROOT)
MAX_INTERVIEW_IMAGES = 2


def _now_local() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _generate_id() -> str:
    return uuid.uuid4().hex[:16]


def _interview_to_dict(interview: DirectorInterview) -> dict:
    images = (
        DirectorInterviewImage.query
        .filter_by(interview_id=interview.id)
        .order_by(DirectorInterviewImage.sort_order)
        .all()
    )
    return {
        "id": interview.id,
        "employeeName": interview.employee_name,
        "employeeId": interview.employee_id,
        "division": interview.division,
        "position": interview.position,
        "interviewDate": interview.interview_date,
        "interviewTopic": interview.interview_topic,
        "interviewContent": json.loads(interview.interview_content or "[]"),
        "conclusion": json.loads(interview.conclusion or "[]"),
        "interviewer": interview.interviewer,
        "interviewee": interview.interviewee,
        "images": [_image_to_dict(img) for img in images],
        "createdAt": interview.created_at,
        "updatedAt": interview.updated_at,
    }


def _image_to_dict(image: DirectorInterviewImage) -> dict:
    return {
        "id": image.id,
        "interviewId": image.interview_id,
        "fileName": image.file_name,
        "filePath": image.file_path,
        "fileSize": image.file_size,
        "mimeType": image.mime_type,
        "sortOrder": image.sort_order,
        "url": f"/api/director-interview-images/{image.file_path}",
    }


# ── CRUD ────────────────────────────────────────────────────────────────────


def list_interviews(year: str = "", month: str = "", keyword: str = "") -> dict:
    query = DirectorInterview.query
    if year:
        query = query.filter(DirectorInterview.interview_date.like(f"{year}-%"))
    if month:
        query = query.filter(DirectorInterview.interview_date.like(f"%-{month}-%"))
    if keyword:
        kw = f"%{keyword}%"
        query = query.filter(
            db.or_(
                DirectorInterview.employee_name.like(kw),
                DirectorInterview.employee_id.like(kw),
                DirectorInterview.interview_topic.like(kw),
            )
        )
    query = query.order_by(
        DirectorInterview.interview_date.desc(), DirectorInterview.id.desc()
    )
    interviews = query.all()
    return {"list": [_interview_to_dict(i) for i in interviews], "total": len(interviews)}


def get_interview(interview_id: str) -> dict | None:
    interview = DirectorInterview.query.get(interview_id)
    if not interview:
        return None
    return _interview_to_dict(interview)


def create_interview(payload: dict) -> dict:
    interview = DirectorInterview(
        id=_generate_id(),
        employee_name=str(payload.get("employeeName", "")).strip(),
        employee_id=str(payload.get("employeeId", "")).strip(),
        division=str(payload.get("division", "")).strip(),
        position=str(payload.get("position", "")).strip(),
        interview_date=str(payload.get("interviewDate", "")).strip(),
        interview_topic=str(payload.get("interviewTopic", "")).strip(),
        interview_content=json.dumps(payload.get("interviewContent", []), ensure_ascii=False),
        conclusion=json.dumps(payload.get("conclusion", []), ensure_ascii=False),
        interviewer=str(payload.get("interviewer", "")).strip(),
        interviewee=str(payload.get("interviewee", "")).strip(),
        created_at=_now_local(),
    )
    db.session.add(interview)
    db.session.commit()
    return _interview_to_dict(interview)


def update_interview(interview_id: str, payload: dict) -> dict | None:
    interview = DirectorInterview.query.get(interview_id)
    if not interview:
        return None

    for key, attr in [
        ("employeeName", "employee_name"),
        ("employeeId", "employee_id"),
        ("division", "division"),
        ("position", "position"),
        ("interviewDate", "interview_date"),
        ("interviewTopic", "interview_topic"),
        ("interviewer", "interviewer"),
        ("interviewee", "interviewee"),
    ]:
        if key in payload:
            setattr(interview, attr, str(payload[key]).strip())

    if "interviewContent" in payload:
        interview.interview_content = json.dumps(payload["interviewContent"], ensure_ascii=False)
    if "conclusion" in payload:
        interview.conclusion = json.dumps(payload["conclusion"], ensure_ascii=False)

    interview.updated_at = _now_local()
    db.session.commit()
    return _interview_to_dict(interview)


def delete_interview(interview_id: str) -> bool:
    interview = DirectorInterview.query.get(interview_id)
    if not interview:
        return False
    images = DirectorInterviewImage.query.filter_by(interview_id=interview_id).all()
    for img in images:
        _delete_image_file(img)
        db.session.delete(img)
    db.session.delete(interview)
    db.session.commit()
    return True


# ── Image management ────────────────────────────────────────────────────────


def _delete_image_file(image: DirectorInterviewImage) -> None:
    try:
        file_path = DIRECTOR_INTERVIEW_IMAGE_ROOT / image.file_path
        if file_path.exists():
            file_path.unlink()
    except Exception:
        pass


def save_interview_images(interview_id: str, files: list[dict]) -> list[dict]:
    interview = DirectorInterview.query.get(interview_id)
    if not interview:
        raise ValueError("约谈记录不存在")

    existing_count = DirectorInterviewImage.query.filter_by(interview_id=interview_id).count()
    if existing_count + len(files) > MAX_INTERVIEW_IMAGES:
        raise ValueError(f"每条约谈记录最多上传 {MAX_INTERVIEW_IMAGES} 张图片")

    current_max_order = (
        db.session.query(db.func.max(DirectorInterviewImage.sort_order))
        .filter_by(interview_id=interview_id)
        .scalar()
    ) or 0

    saved = []
    interview_dir = DIRECTOR_INTERVIEW_IMAGE_ROOT / interview_id
    interview_dir.mkdir(parents=True, exist_ok=True)

    for index, file_info in enumerate(files):
        file_name = str(file_info.get("fileName", "image.png"))
        content_b64 = str(file_info.get("content", ""))
        mime_type = str(file_info.get("mimeType", "image/png"))
        ext = _guess_ext(file_name, mime_type)
        stored_name = f"{uuid.uuid4().hex[:8]}{ext}"
        relative_path = f"{interview_id}/{stored_name}"
        full_path = interview_dir / stored_name

        try:
            raw_bytes = base64.b64decode(content_b64)
        except Exception:
            continue

        full_path.write_bytes(raw_bytes)
        image = DirectorInterviewImage(
            id=_generate_id(),
            interview_id=interview_id,
            file_name=file_name,
            file_path=relative_path,
            file_size=len(raw_bytes),
            mime_type=mime_type,
            sort_order=current_max_order + index + 1,
            created_at=_now_local(),
        )
        db.session.add(image)
        saved.append(image)

    db.session.commit()
    return [_image_to_dict(img) for img in saved]


def delete_interview_image(interview_id: str, image_id: str) -> bool:
    image = DirectorInterviewImage.query.filter_by(
        id=image_id, interview_id=interview_id
    ).first()
    if not image:
        return False
    _delete_image_file(image)
    db.session.delete(image)
    db.session.commit()
    return True


def _guess_ext(file_name: str, mime_type: str) -> str:
    if mime_type:
        ext = mimetypes.guess_extension(mime_type)
        if ext:
            return ext
    return Path(file_name).suffix or ".png"


# ── DOCX Export ─────────────────────────────────────────────────────────────


def export_interview_docx(interview_id: str) -> bytes | None:
    """生成主任座谈会 DOCX（Lampiran D 格式）"""
    interview_data = get_interview(interview_id)
    if not interview_data:
        return None

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")

    def _set_run_east_asian(run, font_name: str):
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _set_cell(table, row: int, col: int, text: str, bold: bool = False, fs: float = 10.5):
        cell = table.cell(row, col)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(fs)
        run.font.name = "Times New Roman"
        return run

    # Parse data
    employee_name = interview_data["employeeName"]
    employee_id = interview_data["employeeId"]
    division = interview_data["division"]
    position = interview_data["position"]
    interview_topic = interview_data["interviewTopic"]
    interview_content = interview_data.get("interviewContent", [])
    conclusion = interview_data.get("conclusion", [])
    interviewer = interview_data.get("interviewer", "")
    interviewee = interview_data.get("interviewee", "")

    try:
        dt = datetime.strptime(interview_data["interviewDate"], "%Y-%m-%d")
        date_display = f"{dt.year} Tahun {dt.month:02d} Bulan {dt.day:02d} Tanggal\n{dt.year} 年 {dt.month:02d} 月 {dt.day:02d} 日"
    except Exception:
        date_display = interview_data["interviewDate"]

    # Build document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lampiran D")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Formulir Catatan Komunikasi PT ITSS Departemen Teknologi Peralatan\n"
        "ITSS公司设备技术部沟通交流记录表"
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # ── Info table: 10 rows × 2 cols ──────────────────────────────────────
    table = doc.add_table(rows=10, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10)

    fields = [
        ("NAMA KARYAWAN / 员工姓名", employee_name),
        ("NO ID / 工号", employee_id),
        ("DIVISI / 车间", division),
        ("JABATAN / 岗位", position),
        ("TANGGAL WAWANCARA / 约谈时间", date_display),
        ("TOPIK WAWANCARA / 约谈主题", interview_topic),
    ]

    for i, (label, value) in enumerate(fields):
        _set_cell(table, i, 0, label, bold=True, fs=10.5)
        _set_cell(table, i, 1, value, fs=10.5)

    # KONTEN WAWANCARA (row 6) — span across both cols for numbered list
    table.cell(6, 0).merge(table.cell(6, 1))
    cell = table.cell(6, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("KONTEN WAWANCARA / 约谈内容:")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    if interview_content:
        for item in interview_content:
            p = cell.add_paragraph()
            run = p.add_run(f"  {item}")
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"

    # Kesimpulan (row 7) — span across both cols
    table.cell(7, 0).merge(table.cell(7, 1))
    cell = table.cell(7, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("Kesimpulan wawancara / 约谈后总结:")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    if conclusion:
        for item in conclusion:
            p = cell.add_paragraph()
            run = p.add_run(f"  {item}")
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"

    # Signatures (rows 8-9): interviewer & interviewee side by side
    table.cell(8, 0).merge(table.cell(9, 0))
    table.cell(8, 1).merge(table.cell(9, 1))

    for col_idx, (label, value) in enumerate([
        (f"TTD PEWAWANCARA\n约谈人签字\n\n{interviewer}", interviewer),
        (f"TTD KARYAWAN\n被约谈人签字\n\n{interviewee}", interviewee),
    ]):
        cell = table.cell(8, col_idx)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Catatan / 备注: Harus diisi dalam 2 bahasa（中印双语记录）\n"
        "Form ini dapat menjadi dasar karyawan bermasalah atau berperforma kurang\n"
        "此表可作为用于针对表现不好、问题员工的约谈依据"
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.italic = True

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
